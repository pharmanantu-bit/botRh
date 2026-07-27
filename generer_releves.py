import calendar
from docx import Document

SOURCE = "documents/Relevé_heures_Mai_26 (1) (1).docx"
ANNEE = 26  # 2026

MOIS = {
    1: "JANVIER", 2: "FEVRIER", 3: "MARS", 4: "AVRIL",
    5: "MAI", 6: "JUIN", 7: "JUILLET", 8: "AOUT",
    9: "SEPTEMBRE", 10: "OCTOBRE", 11: "NOVEMBRE", 12: "DECEMBRE"
}

MOIS_FICHIER = {
    1: "Janvier", 2: "Fevrier", 3: "Mars", 4: "Avril",
    5: "Mai", 6: "Juin", 7: "Juillet", 8: "Aout",
    9: "Septembre", 10: "Octobre", 11: "Novembre", 12: "Decembre"
}


def remplacer_texte(paragraph, ancien, nouveau):
    for run in paragraph.runs:
        if ancien in run.text:
            run.text = run.text.replace(ancien, nouveau)


def generer_releve(mois):
    mois_precedent = 12 if mois == 1 else mois - 1
    annee_prec = ANNEE - 1 if mois == 1 else ANNEE

    nom_mois = MOIS[mois]
    nom_mois_prec = MOIS[mois_precedent]

    nb_jours_prec = calendar.monthrange(2000 + annee_prec, mois_precedent)[1]
    nb_jours_mois = calendar.monthrange(2000 + ANNEE, mois)[1]

    # Période du relevé : du 26 du mois précédent au 25 du mois courant inclus
    # (alignée sur le formulaire en ligne — NB : script legacy, les mails
    # n'attachent plus de document Word).
    jours_prec = list(range(26, nb_jours_prec + 1))
    jours_mois = list(range(1, 25 + 1))

    doc = Document(SOURCE)

    # Remplacer le titre
    for para in doc.paragraphs:
        if "MAI 26" in para.text:
            remplacer_texte(para, "MAI 26", f"{nom_mois} {ANNEE}")

    # Mettre à jour le tableau
    table = doc.tables[0]
    section = None
    idx_prec = 0
    idx_mois = 0

    for row in table.rows:
        texte = row.cells[0].text.strip()

        if texte == "AVRIL":
            section = "prec"
            for cell in row.cells:
                for para in cell.paragraphs:
                    remplacer_texte(para, "AVRIL", nom_mois_prec)

        elif texte == "MAI":
            section = "mois"
            for cell in row.cells:
                for para in cell.paragraphs:
                    remplacer_texte(para, "MAI", nom_mois)

        elif texte.isdigit():
            if section == "prec" and idx_prec < len(jours_prec):
                for para in row.cells[0].paragraphs:
                    remplacer_texte(para, texte, str(jours_prec[idx_prec]))
                idx_prec += 1
            elif section == "mois" and idx_mois < len(jours_mois):
                for para in row.cells[0].paragraphs:
                    remplacer_texte(para, texte, str(jours_mois[idx_mois]))
                idx_mois += 1

    nom_fichier = f"documents/Releve_heures_{MOIS_FICHIER[mois]}_{2000+ANNEE}.docx"
    doc.save(nom_fichier)
    print(f"Cree : {nom_fichier}")


for m in range(1, 13):
    generer_releve(m)

print("\nTous les releves ont ete generes !")
