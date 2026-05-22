import re
from docx import Document
import pdfplumber


def extraire_nombre(texte):
    texte = texte.strip().replace(",", ".").replace("h", ".")
    match = re.search(r"\d+\.?\d*", texte)
    return float(match.group()) if match else 0.0


def parse_releve_docx(filepath):
    doc = Document(filepath)
    heures_plus = None
    heures_moins = None

    for para in doc.paragraphs:
        texte = para.text.strip()
        if "Heures en plus" in texte and "Sous total" in texte:
            val = re.search(r":\s*([\d,. ]+)", texte)
            if val:
                heures_plus = extraire_nombre(val.group(1))
        elif "Heures en moins" in texte and "Sous total" in texte:
            val = re.search(r":\s*([\d,. ]+)", texte)
            if val:
                heures_moins = extraire_nombre(val.group(1))

    if heures_plus is None or heures_moins is None:
        heures_plus, heures_moins = 0.0, 0.0
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                if len(cells) >= 4 and cells[0].isdigit():
                    heures_plus += extraire_nombre(cells[2])
                    heures_moins += extraire_nombre(cells[3])

    return heures_plus or 0.0, heures_moins or 0.0


def parse_releve_pdf(filepath):
    heures_plus = None
    heures_moins = None

    with pdfplumber.open(filepath) as pdf:
        texte_complet = "\n".join(p.extract_text() or "" for p in pdf.pages)

    # Chercher "Soit au total"
    match_total = re.search(r"Soit au total\s*[:\xa0]*\s*([\d,. ]+)\s*h", texte_complet, re.IGNORECASE)
    if match_total:
        heures_plus = extraire_nombre(match_total.group(1))
        heures_moins = 0.0

    # Chercher les sous-totaux
    match_plus = re.search(r"Heures en plus\s*:\s*([\d,. ]+)", texte_complet, re.IGNORECASE)
    match_moins = re.search(r"Heures en moins\s*:\s*([\d,. ]+)", texte_complet, re.IGNORECASE)
    if match_plus:
        heures_plus = extraire_nombre(match_plus.group(1))
    if match_moins:
        heures_moins = extraire_nombre(match_moins.group(1))

    # Fallback : compter les occurrences d'heures dans le tableau
    if heures_plus is None:
        heures = re.findall(r"(\d+)\s*heures?", texte_complet, re.IGNORECASE)
        heures_plus = sum(int(h) for h in heures)
        heures_moins = 0.0

    return heures_plus or 0.0, heures_moins or 0.0


def parse_releve(filepath):
    if filepath.lower().endswith(".pdf"):
        heures_plus, heures_moins = parse_releve_pdf(filepath)
    else:
        heures_plus, heures_moins = parse_releve_docx(filepath)

    return {
        "heures_plus": round(heures_plus, 2),
        "heures_moins": round(heures_moins, 2),
        "solde": round(heures_plus - heures_moins, 2),
    }
