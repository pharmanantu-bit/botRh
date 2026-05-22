from docx import Document
from docx.shared import Pt
import os

doc = Document()
doc.add_heading("Relevé mensuel des heures — Mai 2026", 0)

doc.add_paragraph("Nom : ___________________________")
doc.add_paragraph("Prénom : ___________________________")
doc.add_paragraph("Poste : ___________________________")
doc.add_paragraph("")

table = doc.add_table(rows=1, cols=4)
table.style = "Table Grid"
hdr = table.rows[0].cells
hdr[0].text = "Date"
hdr[1].text = "Heure arrivée"
hdr[2].text = "Heure départ"
hdr[3].text = "Total heures"

for _ in range(20):
    row = table.add_row().cells
    row[0].text = ""
    row[1].text = ""
    row[2].text = ""
    row[3].text = ""

doc.add_paragraph("")
doc.add_paragraph("Signature : ___________________________")

path = os.path.join("documents", "releve_heures_mai_2026.docx")
doc.save(path)
print(f"Document créé : {path}")
